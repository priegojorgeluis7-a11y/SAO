#!/usr/bin/env python3.11
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  PLANTILLA DE REPORTE — Estilo ATTRAPI / Gobierno de México                 ║
║  Sistema SAO — Unidad de Adquisición de Derechos de Vía                     ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  USO:                                                                        ║
║    1. Copia este archivo con el nombre del proyecto:                         ║
║       cp plantilla_reporte_attrapi.py reporte_MIPROYECTO_FECHA.py           ║
║    2. Edita solo la sección ═══ CONFIGURACIÓN ═══ (líneas ~40-120)          ║
║    3. Ejecuta: /opt/homebrew/bin/python3.11 reporte_MIPROYECTO_FECHA.py     ║
║    4. El .docx se guarda en la ruta que definas en OUTPUT_PATH               ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  ═══ CONFIGURACIÓN DEL PROYECTO (EDITA AQUÍ) ═══                           ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

# ─── Metadatos del reporte ────────────────────────────────────────────────────
PROYECTO_NOMBRE     = "Tren Saltillo–Nuevo Laredo (TSNL)"
PROYECTO_CLAVE      = "TSNL"
FECHA_CORTE         = "20 de mayo de 2026"
FECHA_CORTE_CORTA   = "20/05/2026"
REFERENCIA_DOC      = "SAO-TSNL-RPT-2026-0520"
ELABORADO_POR       = "Unidad de Adquisición de Derechos de Vía — ATTRAPI"
CLASIFICACION       = "Informe interno — Uso restringido"
CIUDAD              = "Ciudad de México"
OUTPUT_PATH         = "/Users/jorgeluispriegocruz/Projects/SAO-clean/TSNL_Reporte_Avances_ATTRAPI_20may2026.docx"

# ─── Ruta al logo ATTRAPI (imagen del header — extraída del docx base) ────────
LOGO_PATH = "/tmp/attrapi_assets/word/media/image104.png"

# ─── Resumen ejecutivo — tabla de KPIs ───────────────────────────────────────
# Formato: ('Indicador', 'Valor', 'Detalle')
KPI_ROWS = [
    ('Total de actividades',          '104',   'En todos los frentes'),
    ('Completadas',                   '71 (68%)', 'Con evidencia cargada y revisión aprobada'),
    ('Pendientes',                    '28 (27%)', 'En proceso o por programar'),
    ('Canceladas',                    '5 (5%)',   'Sin efecto por cambios de trazo o propietario'),
    ('Evidencias cargadas',           '158',   'Fotografías y documentos digitalizados'),
    ('COPs firmados',                 '4',     'Contratos de pago a propietarios'),
    ('Anuencias obtenidas',           '≥ 5',   'Firmadas en el período'),
    ('Municipios activos',            '14',    'Coahuila, Nuevo León y Tamaulipas'),
    ('Observaciones HIGH abiertas',   '3',     'Desde 9 de mayo de 2026'),
    ('Riesgo ALTO activo',            '1',     'MAFESA PK 158+800–159+400 (Frente 15)'),
]

RESUMEN_TEXTO = (
    "El presente informe sintetiza el estado de avance del proyecto TSNL al corte del 20 de mayo de 2026. "
    "El proyecto comprende la adquisición de derechos de vía en cuatro frentes de trabajo a lo largo de "
    "309 km de trazo ferroviario en los estados de Coahuila, Nuevo León y Tamaulipas."
)
RESUMEN_CIERRE = (
    "Al cierre del período, el 100% de las actividades completadas cuentan con revisión aprobada en el sistema SAO. "
    "El principal riesgo vigente corresponde al predio MAFESA en el Frente 15 (Segmento 15B), "
    "cuya propuesta de acceso aún no ha sido resuelta por el representante del propietario."
)

# ─── Equipo operativo — tabla de operativos ───────────────────────────────────
# Formato: ('Nombre', 'Frentes asignados', 'No. actividades')
EQUIPO_ROWS = [
    ('Isis Monserrat Medina',      '15, 16-17',     '25'),
    ('Fernando Martinez',          '13-14, 16-17',  '17'),
    ('Alejandro Rojas',            '15, 16-17, 18', '15'),
    ('Aarón Briones',              '16-17',         '13'),
    ('Fernando Villarreal',        '16-17',         '11'),
    ('Luis Andrade',               '13-14',         '9'),
    ('Eliasaf Zapata',             '17-18',         '8'),
    ('Luis Ortiz',                 '13-14',         '6'),
]
EQUIPO_ALERTA = (
    "⚠  Se detectaron 7 actividades PENDIENTES sin operativo asignado. "
    "Se recomienda asignar responsable antes de la semana del 21 de mayo."
)

# ─── Secciones del reporte ────────────────────────────────────────────────────
# Cada sección es un dict con:
#   'titulo'    : str — encabezado H1
#   'subtitulo' : str — línea descriptiva (color gris, opcional)
#   'operativos': str — lista de operativos (color gris, opcional)
#   'subsecciones': lista de dicts con:
#       'titulo'  : str — encabezado H2
#       'items'   : lista de ('texto', 'color') donde color es 'verde'|'naranja'|'rojo'|None
#   'page_break': bool — insertar salto de página después (default True)

SECCIONES = [
    {
        'titulo': '3.  Frente 13 y 14 — Saltillo / Ramos Arizpe / García',
        'subtitulo': 'Coahuila de Zaragoza y sur de Nuevo León  |  PK 0 – 94 km',
        'operativos': 'Operativos: Fernando Martinez, Luis Andrade, Luis Ortiz',
        'subsecciones': [
            {
                'titulo': '3.1  Actividades realizadas',
                'items': [
                    ('Ejido Angostura (PK 9+000–10+002): Múltiples socializaciones sobre tenencia de tierra y afectaciones al corredor ferroviario.', None),
                    ('Panteones PK 19+000: Socialización — atención a inquietudes sobre afectaciones al patrimonio municipal.', None),
                    ('PK 30+000–31+000 — Blvd Isidro López: Caminamiento de marcaje DDV con ATTRAPI y LIDERVIC.', None),
                    ('✅  Estación Ramos Arizpe (PK 35+000): Socialización y FIRMA DE ANUENCIA de afectaciones obtenida.', 'verde'),
                    ('✅  Estación García (PK 94+000): Múltiples socializaciones. Se firmaron anuencias de 3 predios.', 'verde'),
                    ('Recorrido 4 Estaciones: Acompañamiento institucional con ATTRAPI y SEDATU.', None),
                    ('Reunión informativa de Estaciones: Con Arq. Valverde, Ing. Selma Baez (SEDATU), Municipios García y Santa Catarina.', None),
                ],
            },
            {
                'titulo': '3.2  Pendientes',
                'items': [
                    ('4 socializaciones adicionales programadas en Ejido Angostura (semana 21 mayo).', 'naranja'),
                    ('Agendar levantamiento topográfico en Estación García tras obtención de anuencias.', 'naranja'),
                ],
            },
        ],
        'page_break': True,
    },
    {
        'titulo': '4.  Frente 15 — Gral. Escobedo / Salinas Victoria / Santa Catarina',
        'subtitulo': 'Zona Metropolitana de Monterrey  |  PK 136 – 159 km',
        'operativos': 'Operativos: Alejandro Rojas, Isis Monserrat (coordinación)',
        'subsecciones': [
            {
                'titulo': '4.1  Actividades realizadas',
                'items': [
                    ('PK 136+200 — ZICA Las Torres, Gral. Escobedo: Reunión con municipio, SEDATU y SICTNL.', None),
                    ('PK 143+900 — Caminamiento DDV: Verificación de afectaciones; localización de propietarios.', None),
                    ('PK 156+500 — Propietario pendiente de firma: posibilidad de firma próxima.', None),
                    ('✅  PK 156+600 — Humberto Treviño Treviño: FIRMA DE ANUENCIA obtenida.', 'verde'),
                    ('🔴  PK 158+800–159+400 — MAFESA: RIESGO ALTO. Propuesta de acceso sin resolución.', 'rojo'),
                    ('PK 159+100 — Base de Mantenimiento: Caminamiento de verificación con Residente.', None),
                    ('PK 159+500 — Edificios Auxiliares: Reunión de coordinación con ATTRAPI, SEDATU, SICTNL.', None),
                    ('San Nicolás de los Garza: Puente superior municipal podría interferir con viaducto del tren.', 'naranja'),
                    ('4a Mesa de Atención ZM Monterrey: Con García, Santa Catarina, SEDATU, SICTNL, ATTRAPI.', None),
                    ('Santa Catarina — Reunión técnica: Presentación planeación Segmento 15A1 y programa 15B.', None),
                ],
            },
            {
                'titulo': '4.2  Pendientes',
                'items': [
                    ('Reunión con Familia Hausser — propietarios PK 136+400.', 'naranja'),
                    ('Firma anuencia propietario 15B - SOMOHANO (PK 156+500) — sin operativo asignado.', 'naranja'),
                    ('Mesa de Atención Social ZM (21 mayo) — Santa Catarina.', 'naranja'),
                    ('🔴  Resolución de acceso MAFESA — riesgo alto activo.', 'rojo'),
                    ('Seguimiento propietario San Nicolás de los Garza.', 'naranja'),
                ],
            },
        ],
        'page_break': True,
    },
    {
        'titulo': '5.  Frente 16 y 17 — Salinas Victoria / Bustamante / Villaldama / Lampazos',
        'subtitulo': 'PK 159 – 255 km  |  Frente con mayor volumen de actividad en el período',
        'operativos': 'Operativos: Aarón Briones, Isis Monserrat, Fernando Villarreal, Alejandro Rojas',
        'subsecciones': [
            {
                'titulo': '5.1  Salinas Victoria (Segmento 16)',
                'items': [
                    ('Caminamientos de marcaje y BDTs: PK 170+200–172+880 (SNL-SLV-035, 038–042). ~10 tramos completados.', None),
                    ('✅  Avalúo INDAABIN — ACRELUX (PK 163+750–164+200): Acompañamiento completo. Avalúo finalizado.', 'verde'),
                    ('✅  Interpuerto Transparque (PK 170+875): Autorizaron levantamientos.', 'verde'),
                    ('✅  Firma Anexos COP (PK 188+713): Documentos actualizados y firmados.', 'verde'),
                    ('OSA Industrial Park (PK 195+530): Pendiente verificar propuesta alternativa con Obras.', 'naranja'),
                    ('⚠  Caminamiento INDAABIN Salinas Victoria pendiente desde 30 abril. Reagendar urgente.', 'rojo'),
                ],
            },
            {
                'titulo': '5.2  Bustamante (Segmento 16–17)',
                'items': [
                    ('Ejido N.C.P.E. Francisco Villa (PK 231+156–234+000): Caminamiento con SEDATU, RAN, Procuraduría.', None),
                    ('Familia Ugarte (PK 232+116–233+315): Marcaje de estacas y levantamiento de BDTs.', None),
                    ('Propietario PK 226+350: Preocupado por acceso y daño a cerca dentro del DDV.', 'naranja'),
                    ('✅  COPs — Aldo Saul Leal Maldonado (PK 234+026): 3 COPs firmados.', 'verde'),
                ],
            },
            {
                'titulo': '5.3  Villaldama (Segmento 17)',
                'items': [
                    ('Caminamiento PK 222+324: Marcaje, delimitación y levantamiento de BDTs con SEDATU y LIDERVIC.', None),
                ],
            },
            {
                'titulo': '5.4  Lampazos de Naranjo (Segmento 17)',
                'items': [
                    ('Ganadería La Playa (PK 249+118–249+675): Caminamiento con Familia Quijano Dávalos.', None),
                    ('✅  Firma COP Quijano Dávalos (PK 251+645–253+349): COP zona norte firmado. 1,703 m / 10,120 m².', 'verde'),
                ],
            },
            {
                'titulo': '5.5  Pendientes Frente 16 y 17',
                'items': [
                    ('Reunión de negociación — Salinas Victoria (Isis Monserrat, 20–21 mayo).', 'naranja'),
                    ('Marcaje predio Villaldama — sin operativo asignado.', 'naranja'),
                    ('Asamblea corrección área de adquisición — Lampazos (Eliasaf Zapata, PK 285+119).', 'naranja'),
                    ('Resolución PSV PK 195+530 — coordinación con equipo de Obras.', 'naranja'),
                    ('Propuesta de solución acceso propietario PK 226+350 (Bustamante).', 'naranja'),
                ],
            },
        ],
        'page_break': True,
    },
    {
        'titulo': '6.  Frente 18 — Lampazos de Naranjo / Nuevo Laredo',
        'subtitulo': 'PK 259 – 309 km  |  Frente en etapa inicial de acercamiento',
        'operativos': 'Operativos: Eliasaf Zapata, Alejandro Rojas',
        'subsecciones': [
            {
                'titulo': '6.1  Actividades realizadas',
                'items': [
                    ('✅  Presentación ante SICTNL Nuevo Laredo: Papelería y formatos para DDV entregados.', 'verde'),
                    ('✅  Reconocimiento de predios Seg 18 (PK 259+351–309+392) con personal del municipio de Lampazos.', 'verde'),
                ],
            },
            {
                'titulo': '6.2  Pendientes',
                'items': [
                    ('Asamblea corrección área de adquisición — Lampazos (Eliasaf Zapata, PK 285+119).', 'naranja'),
                    ('Planeación del próximo ciclo de acercamientos con propietarios en Nuevo Laredo.', 'naranja'),
                ],
            },
        ],
        'page_break': True,
    },
]

# ─── Actividades canceladas ───────────────────────────────────────────────────
CANCELADAS_TEXTO = (
    "Se registraron 5 actividades con estado CANCELADO en el sistema SAO. Las cancelaciones corresponden "
    "principalmente a cambios de trazo, propietarios ilocalizables o actividades sustituidas por versiones actualizadas."
)
CANCELADAS_ROWS = [
    ('Cambio de trazo o ajuste de PK',                   '2'),
    ('Actividad sustituida por versión actualizada',      '2'),
    ('Propietario ilocalizable / sin efecto',             '1'),
]
CANCELADAS_CIERRE = "Estas actividades no impactan el avance general del proyecto."

# ─── Riesgos ──────────────────────────────────────────────────────────────────
# Formato: ('Nivel', 'Descripción', 'Frente', 'Responsable', 'Acción requerida')
RIESGOS_ROWS = [
    ('🔴 ALTO',   'MAFESA PK 158+800–159+400: propuesta de acceso sin resolución',
     '15', 'Alejandro Rojas / ATTRAPI', 'Definir propuesta y agendar reunión urgente'),
    ('🟡 MEDIO',  'PSV PK 195+530: geometría paso superior vehicular pendiente con Obras',
     '16', 'Aarón Briones', 'Coordinar reunión con equipo Obras'),
    ('🟡 MEDIO',  'Caminamiento INDAABIN Salinas Victoria: pendiente desde 30 abril',
     '16', 'Aarón Briones', 'Reagendar antes de 22 mayo'),
    ('🟡 MEDIO',  'Quijano Dávalos zona sur + remanente: respuesta esperada ~15 días',
     '17', 'Isis Monserrat', 'Dar seguimiento semana 21 mayo'),
    ('🟡 MEDIO',  'Propietario PK 226+350 Bustamante: daño a cerca dentro del DDV',
     '16-17', 'Fernando Villarreal', 'Proponer solución de acceso'),
    ('🟠 INFO',   '3 observaciones HIGH abiertas en SAO desde 9 mayo',
     'Global', 'L. Andrade / A. Briones', 'Cerrar en sistema SAO'),
    ('🟠 INFO',   '7 actividades PENDIENTES sin operativo asignado',
     'Global', 'Coordinación ATTRAPI', 'Asignar operativo inmediato'),
    ('🟠 INFO',   'Proyecto sin fecha de cierre (end_date: null en SAO)',
     'Global', 'Coordinación ATTRAPI', 'Definir fecha estimada de cierre'),
]

# ─── Instituciones ────────────────────────────────────────────────────────────
# Formato: ('Institución', 'Rol', 'Frentes')
INSTITUCIONES_ROWS = [
    ('ATTRAPI',                'Coordinación general adquisición DDV',  'Todos'),
    ('SEDATU',                 'Acompañamiento social y asesoría agraria', '13-14, 15, 16-17'),
    ('SICTNL',                 'Enlace estatal Nuevo León',             '15, 16-17'),
    ('RAN',                    'Registro Agrario Nacional',             '16-17'),
    ('Procuraduría Agraria',   'Asesoría ejidal',                       '16-17'),
    ('INDAABIN',               'Avalúos de predios',                    '16'),
    ('LIDERVIC',               'Levantamientos topográficos',           '13-14, 17'),
    ('Municipio García',       'Coordinación y anuencias',              '13-14'),
    ('Municipio Ramos Arizpe', 'Coordinación y anuencias',              '13-14'),
    ('Municipio Santa Catarina','Mesas de atención ZM',                 '15'),
    ('Municipio Lampazos',     'Reconocimiento predios',                '17-18'),
    ('Municipio Nuevo Laredo', 'Presentación del proyecto',             '18'),
]

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  ═══ LIBRERÍA DE ESTILOS ATTRAPI (no editar) ═══                           ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colores corporativos
GUINDA       = RGBColor(0x6B, 0x18, 0x31)
GUINDA_LIGHT = RGBColor(0x9B, 0x2C, 0x4B)
NEGRO        = RGBColor(0x1A, 0x1A, 0x1A)
GRIS         = RGBColor(0x55, 0x55, 0x55)
VERDE        = RGBColor(0x1B, 0x5E, 0x20)
NARANJA      = RGBColor(0xB7, 0x54, 0x09)
ROJO         = RGBColor(0x8B, 0x00, 0x00)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
GUINDA_HEX   = '6B1831'

COLOR_MAP = {'verde': VERDE, 'naranja': NARANJA, 'rojo': ROJO, None: None}


def _add_border_bottom(paragraph, color_hex, sz='6'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '1');    b.set(qn('w:color'), color_hex)
    pBdr.append(b); pPr.append(pBdr)


def _add_border_top(paragraph, color_hex, sz='6'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    t = OxmlElement('w:top')
    t.set(qn('w:val'), 'single'); t.set(qn('w:sz'), sz)
    t.set(qn('w:space'), '4');    t.set(qn('w:color'), color_hex)
    pBdr.append(t); pPr.append(pBdr)


def _cell_shading(cell, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def build_document():
    doc = Document()

    # Márgenes
    s = doc.sections[0]
    s.page_width   = Inches(8.5); s.page_height   = Inches(11)
    s.left_margin  = Inches(1.18); s.right_margin  = Inches(0.98)
    s.top_margin   = Inches(1.0);  s.bottom_margin = Inches(1.0)
    s.header_distance = Inches(0.35); s.footer_distance = Inches(0.35)

    # Header — logo oficial
    hdr = s.header
    for p in hdr.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    hp = hdr.paragraphs[0]; hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.add_run().add_picture(LOGO_PATH, width=Inches(5.8))
    _add_border_bottom(hp, GUINDA_HEX, sz='12')

    # Footer — referencia + número de página
    ftr = s.footer
    for p in ftr.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    fp = ftr.paragraphs[0]; fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_border_top(fp, GUINDA_HEX)
    r1 = fp.add_run(f"{REFERENCIA_DOC}  |  Corte {FECHA_CORTE_CORTA}  |  Pág. ")
    r1.font.size = Pt(8); r1.font.color.rgb = GRIS; r1.font.name = 'Calibri'
    for tag, val in [('begin', None), (None, 'PAGE'), ('end', None)]:
        r = fp.add_run()
        r.font.size = Pt(8); r.font.color.rgb = GRIS; r.font.name = 'Calibri'
        if tag == 'begin' or tag == 'end':
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), tag)
        else:
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve'); el.text = val
        r._r.append(el)

    # Estilos tipográficos
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'; h1.font.size = Pt(14); h1.font.bold = True
    h1.font.color.rgb = GUINDA
    h1.paragraph_format.space_before = Pt(18); h1.paragraph_format.space_after = Pt(6)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'; h2.font.size = Pt(11); h2.font.bold = True
    h2.font.color.rgb = GUINDA_LIGHT
    h2.paragraph_format.space_before = Pt(10); h2.paragraph_format.space_after = Pt(4)

    nm = doc.styles['Normal']
    nm.font.name = 'Calibri'; nm.font.size = Pt(11); nm.font.color.rgb = NEGRO
    nm.paragraph_format.space_after = Pt(4)

    lb = doc.styles['List Bullet']
    lb.font.name = 'Calibri'; lb.font.size = Pt(10.5)
    lb.paragraph_format.space_after = Pt(2)
    lb.paragraph_format.left_indent = Inches(0.3)

    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    _add_border_bottom(p, GUINDA_HEX)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def body(doc, text, bold=False, color=None, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style='Normal')
    p.alignment = align
    r = p.add_run(text)
    if bold: r.font.bold = True
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p


def bullet(doc, text, color=None, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    if bold: r.font.bold = True
    return p


def spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def kv_table(doc, rows, col_widths=(2.2, 4.5)):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells
        c0.width = Inches(col_widths[0]); c1.width = Inches(col_widths[1])
        r0 = c0.paragraphs[0].add_run(k)
        r0.font.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = GUINDA
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.size = Pt(10)
    return tbl


def data_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    hrow = tbl.rows[0]
    for j, h in enumerate(headers):
        c = hrow.cells[j]; _cell_shading(c, GUINDA_HEX)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = BLANCO
    for i, row_data in enumerate(rows):
        drow = tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            p = drow.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(val)).font.size = Pt(9.5)
    if col_widths:
        for row in tbl.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    return tbl


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  ═══ GENERADOR DEL REPORTE (no editar) ═══                                 ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def generar_reporte():
    doc = build_document()

    # ── PORTADA ──────────────────────────────────────────────────────────────
    spacer(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('REPORTE DE AVANCES')
    r.font.name = 'Calibri'; r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = GUINDA

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Proyecto {PROYECTO_CLAVE} — {PROYECTO_NOMBRE}")
    r2.font.name = 'Calibri'; r2.font.size = Pt(15); r2.font.bold = True
    r2.font.color.rgb = GUINDA_LIGHT

    spacer(doc)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Corte informativo al {FECHA_CORTE}")
    r3.font.name = 'Calibri'; r3.font.size = Pt(11); r3.font.color.rgb = GRIS
    spacer(doc, 2)

    kv_table(doc, [
        ('Referencia',      REFERENCIA_DOC),
        ('Proyecto',        f"{PROYECTO_CLAVE} — {PROYECTO_NOMBRE}"),
        ('Fecha de corte',  FECHA_CORTE),
        ('Elaborado por',   ELABORADO_POR),
        ('Clasificación',   CLASIFICACION),
    ])
    spacer(doc, 2)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"{CIUDAD}, {FECHA_CORTE}")
    r4.font.size = Pt(10); r4.font.color.rgb = GRIS; r4.font.italic = True

    doc.add_page_break()

    # ── 1. RESUMEN EJECUTIVO ─────────────────────────────────────────────────
    h1(doc, '1.  Resumen Ejecutivo')
    body(doc, RESUMEN_TEXTO)
    spacer(doc)
    body(doc, 'Indicadores clave del período:', bold=True)
    spacer(doc)
    data_table(doc, ['Indicador', 'Valor', 'Detalle'], KPI_ROWS,
               col_widths=[2.8, 1.5, 3.0])
    spacer(doc)
    body(doc, RESUMEN_CIERRE)
    doc.add_page_break()

    # ── 2. EQUIPO OPERATIVO ──────────────────────────────────────────────────
    h1(doc, '2.  Equipo Operativo')
    body(doc, 'El siguiente cuadro muestra los operativos activos en el período:')
    spacer(doc)
    data_table(doc, ['Operativo', 'Frente(s)', 'Actividades'], EQUIPO_ROWS,
               col_widths=[3.0, 2.5, 1.5])
    spacer(doc)
    body(doc, EQUIPO_ALERTA, color=NARANJA)
    doc.add_page_break()

    # ── SECCIONES DE FRENTES ─────────────────────────────────────────────────
    for sec in SECCIONES:
        h1(doc, sec['titulo'])
        if sec.get('subtitulo'):
            body(doc, sec['subtitulo'], color=GRIS)
        if sec.get('operativos'):
            body(doc, sec['operativos'], color=GRIS)
        spacer(doc)
        for sub in sec.get('subsecciones', []):
            h2(doc, sub['titulo'])
            for text, color_key in sub['items']:
                clr = COLOR_MAP.get(color_key)
                is_bold = color_key in ('verde', 'rojo')
                bullet(doc, text, color=clr, bold=is_bold)
        if sec.get('page_break', True):
            doc.add_page_break()

    # ── 7. ACTIVIDADES CANCELADAS ────────────────────────────────────────────
    h1(doc, '7.  Actividades Canceladas')
    body(doc, CANCELADAS_TEXTO)
    spacer(doc)
    data_table(doc, ['Motivo de cancelación', 'Cantidad'], CANCELADAS_ROWS,
               col_widths=[5.0, 1.5])
    body(doc, CANCELADAS_CIERRE, color=GRIS)
    doc.add_page_break()

    # ── 8. RIESGOS ───────────────────────────────────────────────────────────
    h1(doc, '8.  Riesgos y Observaciones Identificados')
    body(doc, 'Riesgos activos y observaciones abiertas al cierre del período:')
    spacer(doc)
    data_table(doc, ['Nivel', 'Descripción', 'Frente', 'Responsable', 'Acción requerida'],
               RIESGOS_ROWS, col_widths=[0.8, 2.5, 0.7, 1.5, 2.2])
    doc.add_page_break()

    # ── 9. INSTITUCIONES ─────────────────────────────────────────────────────
    h1(doc, '9.  Instituciones Participantes en el Período')
    body(doc, 'Instituciones que participaron en coordinación, verificación o acuerdo:')
    spacer(doc)
    data_table(doc, ['Institución / Empresa', 'Rol', 'Frentes'], INSTITUCIONES_ROWS,
               col_widths=[2.3, 2.8, 1.5])
    spacer(doc, 2)

    # Pie final
    body(doc, '─' * 72, color=GRIS)
    body(doc, f"Documento generado por el sistema SAO — {ELABORADO_POR}", color=GRIS, size=9)
    body(doc, f"Fecha de generación: {FECHA_CORTE_CORTA}  |  Referencia: {REFERENCIA_DOC}", color=GRIS, size=9)

    doc.save(OUTPUT_PATH)
    print(f"✅  {OUTPUT_PATH}")
    print(f"   Párrafos: {len(doc.paragraphs)}")


if __name__ == '__main__':
    generar_reporte()
